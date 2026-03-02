"""Document reconstruction service supporting PDF and DOCX output."""

import base64
import io
import logging
from pathlib import Path
from typing import Literal

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

from src.models.document import ExtractedDocument, TranslatedChunk

logger = logging.getLogger(__name__)

OutputFormat = Literal["pdf", "docx"]


class DocumentRebuilder:
    """Rebuild translated documents in PDF or DOCX format."""

    def __init__(self) -> None:
        """Initialize the document rebuilder."""
        template_dir = Path(__file__).parent.parent / "templates"
        self.env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=True,
        )

    def rebuild(
        self,
        document: ExtractedDocument,
        translated_chunks: list[TranslatedChunk],
        output_path: str | Path,
        target_language: str = "pt-BR",
        output_format: OutputFormat = "pdf",
    ) -> Path:
        """
        Rebuild the document with translated content.

        Args:
            document: Original extracted document with structure
            translated_chunks: Translated text chunks
            output_path: Path to save the output file
            target_language: Target language code
            output_format: Output format ("pdf" or "docx")

        Returns:
            Path to the generated file
        """
        output_path = Path(output_path)
        
        if output_format == "docx":
            output_path = output_path.with_suffix(".docx")
        else:
            output_path = output_path.with_suffix(".pdf")
            
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Rebuilding {output_format.upper()} with {len(translated_chunks)} translated blocks...")

        translation_map = {chunk.id: chunk.translated_text for chunk in translated_chunks}
        blocks = self._prepare_blocks(document, translation_map)

        if output_format == "docx":
            self._generate_docx(blocks, document, output_path)
        else:
            page_width, page_height = self._get_page_dimensions(document)
            html_content = self._render_html(
                blocks=blocks,
                page_width=page_width,
                page_height=page_height,
                language=target_language,
            )
            self._generate_pdf(html_content, output_path)

        logger.info(f"{output_format.upper()} successfully rebuilt: {output_path}")
        return output_path

    def _prepare_blocks(
        self,
        document: ExtractedDocument,
        translation_map: dict[str, str],
    ) -> list[dict]:
        """Prepare blocks for rendering, merging text and images."""
        blocks = []

        for text_block in document.blocks:
            translated_text = translation_map.get(text_block.id, text_block.text)

            blocks.append({
                "type": "text",
                "block_type": text_block.block_type,
                "text": translated_text,
                "page_number": text_block.page_number,
                "reading_order": text_block.reading_order,
            })

        for image in document.images:
            blocks.append({
                "type": "image",
                "block_type": "image",
                "image_data": image.image_data,
                "image_format": image.image_format,
                "alt_text": image.alt_text,
                "page_number": image.page_number,
                "reading_order": image.reading_order,
            })

        blocks.sort(key=lambda b: (b["page_number"], b["reading_order"]))

        return blocks

    def _get_page_dimensions(
        self,
        document: ExtractedDocument,
    ) -> tuple[float, float]:
        """Get the most common page dimensions or default to A4."""
        if not document.page_dimensions:
            return (210.0, 297.0)

        dims = list(document.page_dimensions.values())
        if dims:
            return dims[0]

        return (210.0, 297.0)

    def _render_html(
        self,
        blocks: list[dict],
        page_width: float,
        page_height: float,
        language: str,
    ) -> str:
        """Render blocks to HTML using Jinja2 template."""
        template = self.env.get_template("document.html")

        return template.render(
            blocks=blocks,
            page_width=page_width,
            page_height=page_height,
            language=language,
        )

    def _generate_pdf(self, html_content: str, output_path: Path) -> None:
        """Generate PDF from HTML using WeasyPrint."""
        html = HTML(string=html_content)
        html.write_pdf(str(output_path))

    def _generate_docx(
        self,
        blocks: list[dict],
        document: ExtractedDocument,
        output_path: Path,
    ) -> None:
        """Generate DOCX file using python-docx."""
        doc = DocxDocument()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        for block in blocks:
            if block["type"] == "image":
                try:
                    image_data = base64.b64decode(block["image_data"])
                    image_stream = io.BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to add image to DOCX: {e}")
                continue

            text = block["text"]
            block_type = block["block_type"]

            if block_type == "title":
                paragraph = doc.add_heading(text, level=0)
            elif block_type == "heading":
                paragraph = doc.add_heading(text, level=1)
            elif block_type == "section_header":
                paragraph = doc.add_heading(text, level=2)
            elif block_type == "list_item":
                paragraph = doc.add_paragraph(text, style="List Bullet")
            elif block_type == "caption":
                paragraph = doc.add_paragraph(text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.italic = True
            else:
                paragraph = doc.add_paragraph(text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(str(output_path))


# Keep backwards compatibility
PDFRebuilder = DocumentRebuilder


def rebuild_document(
    document: ExtractedDocument,
    translated_chunks: list[TranslatedChunk],
    output_path: str | Path,
    target_language: str = "pt-BR",
    output_format: OutputFormat = "pdf",
) -> Path:
    """
    Convenience function to rebuild a document.

    Args:
        document: Original extracted document
        translated_chunks: Translated chunks
        output_path: Output file path
        target_language: Target language code
        output_format: Output format ("pdf" or "docx")

    Returns:
        Path to generated file
    """
    rebuilder = DocumentRebuilder()
    return rebuilder.rebuild(
        document=document,
        translated_chunks=translated_chunks,
        output_path=output_path,
        target_language=target_language,
        output_format=output_format,
    )


def rebuild_pdf(
    document: ExtractedDocument,
    translated_chunks: list[TranslatedChunk],
    output_path: str | Path,
    target_language: str = "pt-BR",
) -> Path:
    """Backwards compatible function for PDF rebuilding."""
    return rebuild_document(
        document=document,
        translated_chunks=translated_chunks,
        output_path=output_path,
        target_language=target_language,
        output_format="pdf",
    )
